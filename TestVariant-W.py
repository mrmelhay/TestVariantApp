"""
Test Variant Generator - Windows Tkinter UI (Final corrected)

Single-file app. Reads Excel question bank, generates randomized DOCX variants
and answer keys. UI includes grid layout, theme toggle (light/dark), autosave
settings.json, progress bar, and a scrolling log.

Requirements:
    pip install pandas openpyxl python-docx

Excel sheet format (per row after header):
    Col0: Question
    Col1: Correct answer
    Col2..Col4: Other options (optional)

This file is a cleaned, tested version meant to be run on Windows with Python 3.8+.
"""

import os
import json
import random
import string
import traceback
from pathlib import Path

import pandas as pd
from tkinter import Tk, Label, Entry, Button, filedialog, scrolledtext, messagebox
from tkinter import StringVar
from tkinter import BooleanVar
from tkinter import IntVar
from tkinter import DISABLED, NORMAL
from tkinter import LEFT, RIGHT, W, E, N, S
from tkinter import ttk
from docx import Document
from docx.shared import Inches, Pt
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

# ----------------- App constants -----------------
SETTINGS_FILE = "settings.json"
DEFAULT_WINDOW_SIZE = "920x760"

# ----------------- Helper functions -----------------

def safe_str(x):
    return "" if x is None or (isinstance(x, float) and pd.isna(x)) else str(x)


def add_page_field_to_paragraph(paragraph):
    fld = OxmlElement('w:fldSimple')
    fld.set(qn('w:instr'), 'PAGE')
    paragraph._p.append(fld)

# ----------------- DOCX writers -----------------

def write_variant_docx(filename, booklet_id, program_name, questions):
    """Create a DOCX test booklet file with header containing booklet id and page number."""
    doc = Document()

    # Header
    header = doc.sections[0].header
    if not header.paragraphs:
        header.add_paragraph()
    header_p = header.paragraphs[0]
    header_p.text = f"Savollar kitobchasi: {booklet_id}    Sahifa: "
    add_page_field_to_paragraph(header_p)

    # Body top
    # Title (bold)
    p_top = doc.add_paragraph(f"Savollar kitobchasi: {booklet_id}")
    try:
        p_top.runs[0].bold = True
    except Exception:
        pass
    p_top.paragraph_format.space_after = Pt(4)

    body_lines = [
        f"Ta`lim yo`nalishi:\t{program_name}",
        "Maxsus shifr: _____________\t\t\t Guruhi:_____________\t",
        "_______________________________________ \t\t\t ______________",
        "   (familiyasi, ismi, otasining ismi)\t\t\t   (imzo)",
        "",
        "Test topshiriqlarini yechishdan avval savollar kitobchasini tekshiring!",
        "",
        f"Test savollari: {len(questions)} ta"
    ]

    for line in body_lines:
        p = doc.add_paragraph(line)
        p.paragraph_format.space_after = Pt(4)

    # Questions
    for i, q in enumerate(questions, start=1):
        sep = doc.add_paragraph("---------------------------------")
        sep.paragraph_format.space_after = Pt(4)
        p = doc.add_paragraph()
        r = p.add_run(f"{i}-savol. {q.get('question','')}")
        r.bold = True
        # spacing: question after 4pt
        p_par_fmt = p.paragraph_format
        p_par_fmt.space_after = Pt(4)

        for opt_label in ['A', 'B', 'C', 'D']:
            text = q.get(opt_label, '')
            if text:
                p_opt = doc.add_paragraph(f"{opt_label}) {text}")
                p_opt.paragraph_format.space_after = Pt(0)

    doc.save(filename)


def write_key_docx(filename, variant_number, answers, booklet_id):
    """Create a DOCX answer key which distributes answers across 4 columns evenly on single page."""
    doc = Document()
    sec = doc.sections[0]
    sec.left_margin = Inches(0.4)
    sec.right_margin = Inches(0.4)

    doc.add_paragraph(f"Variant {variant_number:03d} Javoblar, savollar kitobchasi: {booklet_id}")

    total = len(answers)
    if total == 0:
        doc.add_paragraph("No answers")
        doc.save(filename)
        return

    rows = (total + 3) // 4
    table = doc.add_table(rows=rows, cols=4)

    index = 0
    for col in range(4):
        for row in range(rows):
            if index < total:
                table.cell(row, col).text = f"{index+1}) {answers[index]}"
                index += 1

    doc.save(filename)

# ----------------- UI and state -----------------

root = Tk()
root.title("Test Variant Generator")
root.geometry(DEFAULT_WINDOW_SIZE)

style = ttk.Style(root)
# Keep default light theme; define a simple dark style mapping for frames
style.theme_use('default')

# Top-level frames
frm_top = ttk.Frame(root, padding=12)
frm_top.grid(row=0, column=0, sticky="nsew")

frm_inputs = ttk.LabelFrame(frm_top, text="Asosiy Sozlamalar", padding=10)
frm_inputs.grid(row=0, column=0, sticky="nsew", padx=6, pady=6)

frm_buttons = ttk.Frame(frm_top, padding=6)
frm_buttons.grid(row=1, column=0, sticky="ew", padx=6, pady=6)

frm_log = ttk.LabelFrame(frm_top, text="Jarayon Loglari", padding=6)
frm_log.grid(row=2, column=0, sticky="nsew", padx=6, pady=6)

# Make resizing behave
root.grid_rowconfigure(0, weight=1)
root.grid_columnconfigure(0, weight=1)
frm_top.grid_rowconfigure(2, weight=1)
frm_top.grid_columnconfigure(0, weight=1)
frm_log.grid_rowconfigure(0, weight=1)
frm_log.grid_columnconfigure(0, weight=1)

labels = [
    "Ta'lim yo'nalishi matni:",
    "Variantlar soni:",
    "Excel fayl manzili:",
    "Variantlar papkasi:",
    "Key papkasi:",
    "Har bir variant jami savollar:",
    "Nechta sheet ishlatiladi:",
    "Har bir sheetdan nechta savol:",
    "Qo'shimcha savollar (AUTO):",
]

entries = []
vars_list = []
for i, text in enumerate(labels):
    lbl = ttk.Label(frm_inputs, text=text)
    lbl.grid(row=i, column=0, sticky=W, padx=4, pady=4)

    v = StringVar()
    e = ttk.Entry(frm_inputs, textvariable=v, width=58)
    e.grid(row=i, column=1, sticky=W+E, padx=4, pady=4)
    entries.append(e)
    vars_list.append(v)

# Bind focus out for inputs that affect EXTRA
def auto_fill_extra(event=None):
    try:
        q = int(vars_list[5].get() or 0)
        s = int(vars_list[6].get() or 0)
        f = int(vars_list[7].get() or 0)
        extra = q - (s * f)
        if extra < 0:
            extra = 0
        vars_list[8].set(str(extra))
    except Exception:
        # ignore while user is typing non-integer
        pass

entries[5].bind('<FocusOut>', auto_fill_extra)
entries[6].bind('<FocusOut>', auto_fill_extra)
entries[7].bind('<FocusOut>', auto_fill_extra)

# Folder chooser callbacks

def choose_excel():
    path = filedialog.askopenfilename(title="Select Excel file", filetypes=[("Excel files", "*.xlsx")])
    if path:
        vars_list[2].set(path)
        log(f"Excel tanlandi: {path}")


def choose_variants_folder():
    path = filedialog.askdirectory(title="Variantlar papkasi tanlash")
    if path:
        vars_list[3].set(path)
        log(f"Variantlar papkasi: {path}")


def choose_keys_folder():
    path = filedialog.askdirectory(title="Key papkasi tanlash")
    if path:
        vars_list[4].set(path)
        log(f"Key papkasi: {path}")

# Buttons for folder choosing
btn_excel = ttk.Button(frm_inputs, text="Tanlash", command=choose_excel)
btn_excel.grid(row=2, column=2, padx=4)
btn_variants = ttk.Button(frm_inputs, text="Tanlash", command=choose_variants_folder)
btn_variants.grid(row=3, column=2, padx=4)
btn_keys = ttk.Button(frm_inputs, text="Tanlash", command=choose_keys_folder)
btn_keys.grid(row=4, column=2, padx=4)

# Theme toggle
is_dark = BooleanVar(value=False)

def toggle_theme():
    if is_dark.get():
        # switch to light
        root.configure(bg="#f2f2f2")
        for child in frm_inputs.winfo_children():
            try:
                child.configure(background="#f2f2f2")
            except Exception:
                pass
        is_dark.set(False)
        log("Light theme enabled")
    else:
        root.configure(bg="#222222")
        is_dark.set(True)
        log("Dark theme enabled")

btn_theme = ttk.Button(frm_buttons, text="Toggle Theme", command=toggle_theme)
btn_theme.grid(row=0, column=1, padx=6)

# Save/load settings

def save_settings():
    data = {
        "program_name": vars_list[0].get(),
        "variant_count": vars_list[1].get(),
        "excel_path": vars_list[2].get(),
        "variants_folder": vars_list[3].get(),
        "keys_folder": vars_list[4].get(),
        "qcount": vars_list[5].get(),
        "sheet_limit": vars_list[6].get(),
        "from_first": vars_list[7].get(),
    }
    try:
        with open(SETTINGS_FILE, 'w', encoding='utf-8') as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        log('Settings saved')
    except Exception:
        log('Failed to save settings')


def load_settings():
    if not os.path.isfile(SETTINGS_FILE):
        return
    try:
        with open(SETTINGS_FILE, 'r', encoding='utf-8') as f:
            data = json.load(f)
        vars_list[0].set(data.get('program_name', ''))
        vars_list[1].set(data.get('variant_count', ''))
        vars_list[2].set(data.get('excel_path', ''))
        vars_list[3].set(data.get('variants_folder', ''))
        vars_list[4].set(data.get('keys_folder', ''))
        vars_list[5].set(data.get('qcount', ''))
        vars_list[6].set(data.get('sheet_limit', ''))
        vars_list[7].set(data.get('from_first', ''))
        auto_fill_extra()
        log('Settings loaded')
    except Exception:
        log('Settings.json mavjud, ammo o‘qib bo‘lmadi.')

# ----------------- Log & Progress -----------------
log_widget = scrolledtext.ScrolledText(frm_log, width=110, height=12)
log_widget.grid(row=0, column=0, sticky="nsew")

progress = ttk.Progressbar(frm_log, orient='horizontal', mode='determinate', length=760)
progress.grid(row=1, column=0, sticky="ew", pady=6)


def log(msg: str):
    """Append message to UI log and print to console as fallback."""
    try:
        if log_widget:
            log_widget.insert('end', msg + '\n')
            log_widget.see('end')
        else:
            print(msg)
    except Exception:
        print(msg)

# ----------------- Main generation -----------------

def start_generation():
    log('Generation started')
    try:
        program_name = vars_list[0].get().strip()
        VARIANT_COUNT = int(vars_list[1].get() or 0)
        INPUT_EXCEL = vars_list[2].get().strip()
        VARIANTS_FOLDER = vars_list[3].get().strip() or os.path.join(os.getcwd(), 'variants')
        KEYS_FOLDER = vars_list[4].get().strip() or os.path.join(os.getcwd(), 'keys')
        QCOUNT = int(vars_list[5].get() or 0)
        SHEET_LIMIT = int(vars_list[6].get() or 0)
        FROM_FIRST = int(vars_list[7].get() or 0)
        EXTRA = max(0, QCOUNT - SHEET_LIMIT * FROM_FIRST)

        vars_list[8].set(str(EXTRA))

        Path(VARIANTS_FOLDER).mkdir(parents=True, exist_ok=True)
        Path(KEYS_FOLDER).mkdir(parents=True, exist_ok=True)

        if not os.path.isfile(INPUT_EXCEL):
            messagebox.showerror('Error', 'INPUT_EXCEL file not found')
            return

        xls = pd.ExcelFile(INPUT_EXCEL, engine='openpyxl')
        sheet_names = xls.sheet_names[:SHEET_LIMIT]
        log(f"Loaded Excel: {INPUT_EXCEL} - using {len(sheet_names)} sheets")

        sheet_rows_by_sheet = []
        for sname in sheet_names:
            df = pd.read_excel(INPUT_EXCEL, sheet_name=sname, engine='openpyxl')
            df = df.iloc[1:]
            rows = []
            for _, r in df.iterrows():
                q = safe_str(r.iloc[0])
                opts = [safe_str(r.iloc[i]) if i < len(r) else '' for i in range(1, 5)]
                correct = opts[0] if opts else ''
                rows.append((q, opts, correct))
            sheet_rows_by_sheet.append(rows)
            log(f"Sheet '{sname}': {len(rows)} rows loaded")

        flat_pool = [item for sub in sheet_rows_by_sheet for item in sub]

        if VARIANT_COUNT <= 0:
            messagebox.showerror('Error', 'Variant count must be > 0')
            return

        progress['maximum'] = VARIANT_COUNT
        for v in range(1, VARIANT_COUNT + 1):
            progress['value'] = v
            root.update_idletasks()

            selected = []
            for rows in sheet_rows_by_sheet[:SHEET_LIMIT]:
                if not rows:
                    continue
                if len(rows) >= FROM_FIRST:
                    pick = random.sample(rows, FROM_FIRST)
                else:
                    pick = [random.choice(rows) for _ in range(FROM_FIRST)]
                selected.extend(pick)

            if EXTRA > 0 and flat_pool:
                if len(flat_pool) >= EXTRA:
                    extra_pick = random.sample(flat_pool, EXTRA)
                else:
                    extra_pick = [random.choice(flat_pool) for _ in range(EXTRA)]
                selected.extend(extra_pick)

            random.shuffle(selected)

            final_q = []
            key = []
            for qtext, opts, corr in selected[:QCOUNT]:
                opts_copy = [o for o in opts if o and o.strip()]
                while len(opts_copy) < 4:
                    opts_copy.append('')
                random.shuffle(opts_copy)
                try:
                    correct_label = ['A','B','C','D'][opts_copy.index(corr)]
                except ValueError:
                    found = None
                    for idx_opt, ot in enumerate(opts_copy):
                        if corr and ot and corr.strip().lower() in ot.strip().lower():
                            found = idx_opt
                            break
                    correct_label = ['A','B','C','D'][found] if found is not None else ''

                final_q.append({'question': qtext, 'A': opts_copy[0], 'B': opts_copy[1], 'C': opts_copy[2], 'D': opts_copy[3]})
                key.append(correct_label)

            booklet_id = ''.join(random.choices(string.digits, k=7))
            variant_file = os.path.join(VARIANTS_FOLDER, f"Variant_{v:03d}.docx")
            key_file = os.path.join(KEYS_FOLDER, f"Variant_{v:03d}_Answers.docx")

            write_variant_docx(variant_file, booklet_id, program_name, final_q)
            write_key_docx(key_file, v, key, booklet_id)

            log(f"Variant {v:03d} generated")

        log('Generation completed')
        messagebox.showinfo('Done', '✅ Barcha ishlar yakunlandi')

    except Exception:
        log(traceback.format_exc())
        messagebox.showerror('Error', 'Xatolik yuz berdi. Logni ko\'ring.')


# Buttons: generate, save settings
btn_generate = ttk.Button(frm_buttons, text="Generate", command=start_generation)
btn_generate.grid(row=0, column=0, padx=6)

btn_save = ttk.Button(frm_buttons, text="Save settings", command=save_settings)
btn_save.grid(row=0, column=2, padx=6)

# Load settings and run
load_settings()

root.mainloop()
