"""
Test Variant Generator
Generates N variants (DOCX) and answer keys (DOCX) from an Excel question bank.

Configurable variables at top.

Requires:
    pip install pandas openpyxl python-docx

Input Excel expectations (flexible):
    - Each sheet contains rows where one column is question text and four columns are options.
    - A column indicates the correct answer (letter A/B/C/D) or the correct option text.

If your Excel uses different column names, update the COLUMN_ALIASES mapping.
"""
# ============================
import os
import random
import string
import pandas as pd
from docx import Document
from pathlib import Path


# ========== CONFIG ==========
VARIANT_COUNT = 140
INPUT_EXCEL = r"D:\MyProjects\PythonProject\TestVariantApp\test-base\60111800-Xorijiy til va adabiyoti (ingliz tili).xlsx"
ASOSIY_FOLDER = r"D:\MyProjects\PythonProject\TestVariantApp"
VARIANTS_FOLDER = os.path.join(ASOSIY_FOLDER, "variants")
KEYS_FOLDER = os.path.join(ASOSIY_FOLDER, "keys")
FROM_FIRST_8_PER_SHEET = 6
EXTRA_FROM_OTHER_SHEETS = 2
QUESTIONS_PER_VARIANT = 50
BOOKLET_ID_LENGTH = 7
COLUMN_ALIASES = {
    'question': ['savol', 'question', 'text'],
    'correct': ['to\'g\'ri javob', 'correct', 'answer'],
}
# Create output directories
Path(VARIANTS_FOLDER).mkdir(parents=True, exist_ok=True)
Path(KEYS_FOLDER).mkdir(parents=True, exist_ok=True)

# Helpers

def find_column(df, aliases):
    """Return the first matching column name in df for the aliases list, or None."""
    cols = [c.lower().strip() for c in df.columns]
    for a in aliases:
        a_norm = a.lower().strip()
        for orig, norm in zip(df.columns, cols):
            if a_norm == norm:
                return orig
    return None


def detect_columns(df):
    mapping = {}
    cols = list(df.columns)
    # 1st column = question
    mapping['question'] = cols[0]
    # 2nd column = correct answer
    mapping['correct'] = cols[1] if len(cols) > 1 else None
    # 3-5 columns = alternative options
    mapping['option_a'] = cols[1] if len(cols) > 1 else None  # correct answer treated as option A
    mapping['option_b'] = cols[2] if len(cols) > 2 else None
    mapping['option_c'] = cols[3] if len(cols) > 3 else None
    mapping['option_d'] = cols[4] if len(cols) > 4 else None
    return mapping


def read_sheet_rows(df, colmap):
    rows = []
    for idx, r in df.iterrows():
        q = r[colmap['question']] if colmap['question'] in r and pd.notna(r[colmap['question']]) else None
        a = r[colmap['option_a']] if colmap['option_a'] in r and pd.notna(r[colmap['option_a']]) else None
        b = r[colmap['option_b']] if colmap['option_b'] in r and pd.notna(r[colmap['option_b']]) else None
        c = r[colmap['option_c']] if colmap['option_c'] in r and pd.notna(r[colmap['option_c']]) else None
        d = r[colmap['option_d']] if colmap['option_d'] in r and pd.notna(r[colmap['option_d']]) else None
        corr = r[colmap['correct']] if (colmap['correct'] in r and pd.notna(r[colmap['correct']])) else None
        # require question and at least two options
        if q and any([a, b, c, d]):
            rows.append({'question': str(q), 'A': str(a) if a and not pd.isna(a) else '',
                         'B': str(b) if b and not pd.isna(b) else '',
                         'C': str(c) if c and not pd.isna(c) else '',
                         'D': str(d) if d and not pd.isna(d) else '',
                         'correct': str(corr) if corr and not pd.isna(corr) else ''})
    return rows


def normalize_correct_label(correct, row):
    """Try to convert 'correct' into one of 'A','B','C','D'. If correct is an option text, match it."""
    if not correct:
        return None
    corr = str(correct).strip()
    # if single letter
    if len(corr) == 1 and corr.upper() in ['A','B','C','D']:
        return corr.upper()
    # sometimes '1'.. or full text; try matching to option values
    for label in ['A','B','C','D']:
        opt = (row.get(label) or '').strip()
        if not opt:
            continue
        # compare lowercased normalized
        if corr.lower() == opt.lower():
            return label
    # fallback: if corr appears somewhere in option text
    for label in ['A','B','C','D']:
        opt = (row.get(label) or '').strip()
        if opt and corr.lower() in opt.lower():
            return label
    return None


def shuffle_question(row):
    """Shuffle options and return shuffled row plus new_correct_label"""
    options = [('A', row['A']), ('B', row['B']), ('C', row['C']), ('D', row['D'])]
    # remove empty options but keep labels to preserve mapping; if fewer than 4, still shuffle present ones
    non_empty = [o for o in options if o[1].strip() != '']
    if len(non_empty) < 2:
        # nothing to shuffle
        return row, None
    random.shuffle(non_empty)
    # reassign labels A-D to shuffled options
    new_labels = ['A','B','C','D']
    new_row = {'question': row['question']}
    for i, opt in enumerate(non_empty):
        new_row[new_labels[i]] = opt[1]
    # fill remaining with empty strings
    for j in range(len(non_empty), 4):
        new_row[new_labels[j]] = ''
    # find which option text was correct originally
    orig_label = row.get('correct_raw')
    orig_correct_label = row.get('correct')
    # prefer orig_correct_label if present
    if orig_correct_label and orig_correct_label in ['A','B','C','D']:
        orig_text = row.get(orig_correct_label)
    else:
        # try orig_label text matching
        orig_text = orig_label
    # now find new label
    new_correct = None
    if orig_text:
        for label in ['A','B','C','D']:
            if new_row.get(label) and new_row[label].strip().lower() == str(orig_text).strip().lower():
                new_correct = label
                break
        if not new_correct:
            # partial match
            for label in ['A','B','C','D']:
                if new_row.get(label) and str(orig_text).strip().lower() in new_row[label].strip().lower():
                    new_correct = label
                    break
    return new_row, new_correct


def write_variant_docx(filename, booklet_id, program_name, group_name, questions):
    doc = Document()
    # simple styling
    p = doc.add_paragraph()
    run = p.add_run(f"Savollar kitobchasi:   {booklet_id}")
    run.bold = True
    doc.add_paragraph(f"Ta`lim yo`nalishi: \tMatematika va informatika (kunduzgi)")
    doc.add_paragraph("Maxsus shifr:\t")
    doc.add_paragraph("Guruhi:\t")
    doc.add_paragraph("")
    doc.add_paragraph("(familiyasi, ismi, otasining ismi)\t\t\t(imzo)")
    doc.add_paragraph("")
    doc.add_paragraph("Test topshiriqlarini yechishdan avval savollar kitobchasini varaqlab savollar soni to‘liq mavjudligini tekshiring. Agar savollar soni kamligi aniqlansa yoki savollar kitobchasida muammo bo‘lsa imtihon nazoratchisiga ma’lum qiling!")
    doc.add_paragraph("")
    doc.add_paragraph(f"Test savollari: {len(questions)} ta")
    doc.add_paragraph("")
    # questions
    for i, q in enumerate(questions, start=1):
        # separator line
        doc.add_paragraph("---------------------------------")
        # bold question
        p_q = doc.add_paragraph()
        run_q = p_q.add_run(f"{i}-savol. {q['question']}")
        run_q.bold = True
        # options A-D
        for opt_label in ['A','B','C','D']:
            text = q.get(opt_label) or ''
            if text.strip():
                doc.add_paragraph(f"{opt_label}) {text}")
    doc.save(filename)


def write_key_docx(filename, variant_number, answers):
    doc = Document()
    doc.add_paragraph(f"Variant {variant_number:03d} javoblari:")
    doc.add_paragraph("")
    for i, ans in enumerate(answers, start=1):
        doc.add_paragraph(f"{i}) {ans}")
    # ensure single page by default; save
    doc.save(filename)


# ========================
# Main processing
# ========================

print("Reading Excel file:", INPUT_EXCEL)
xls = pd.ExcelFile(INPUT_EXCEL, engine='openpyxl')
all_sheets = xls.sheet_names
print("Sheets found:", all_sheets)

# Pre-read sheets into lists of question rows
sheet_rows = {}
for sname in all_sheets:
    df = pd.read_excel(xls, sheet_name=sname, engine='openpyxl')
    # Skip header row (first row)
    df = df.iloc[1:]  # Skip header row
    colmap = detect_columns(df)
    # if some option columns missing, we still attempt using position-based fallback
    # fallback: try positional mapping if detection failed
    if not colmap['question']:
        # assume first column
        colmap['question'] = df.columns[0]
    # attempt to fill missing option columns by position
    cols = list(df.columns)
    # find indices for question
    q_idx = cols.index(colmap['question'])
    # try to map next 4 columns as options if explicit not found
    for i, opt_key in enumerate(['option_a','option_b','option_c','option_d']):
        if not colmap[opt_key]:
            pos = q_idx + 1 + i
            if pos < len(cols):
                colmap[opt_key] = cols[pos]
    # correct
    if not colmap['correct']:
        # try last column
        colmap['correct'] = cols[-1]
    rows = read_sheet_rows(df, colmap)
    # normalize correct labels
    for r in rows:
        r['correct_raw'] = r['correct']
        r['correct'] = normalize_correct_label(r['correct_raw'], r)
    sheet_rows[sname] = rows
    print(f"Sheet '{sname}': {len(rows)} usable rows detected")

# Validate that first 8 sheets exist
if len(all_sheets) < 8:
    raise SystemExit("Excel must contain at least 8 sheets as per spec.")

other_sheet_names = all_sheets[8:]

# For reproducibility optionally set a seed
# random.seed(42)

for variant_num in range(1, VARIANT_COUNT + 1):
    # collect questions
    selected = []
    # from sheets 1..8 (by order in workbook)
    for idx in range(8):
        sname = all_sheets[idx]
        rows = sheet_rows[sname]
        if len(rows) >= FROM_FIRST_8_PER_SHEET:
            picked = random.sample(rows, FROM_FIRST_8_PER_SHEET)
        else:
            # sample with replacement if not enough
            picked = [random.choice(rows) for _ in range(FROM_FIRST_8_PER_SHEET)]
        # deep-copy minimal fields
        for p in picked:
            selected.append({'question': p['question'], 'A': p['A'], 'B': p['B'], 'C': p['C'], 'D': p['D'], 'correct': p['correct'], 'correct_raw': p['correct_raw']})
    # extras from other sheets
    pool = []
    for sname in other_sheet_names:
        pool.extend(sheet_rows[sname])
    if len(pool) >= EXTRA_FROM_OTHER_SHEETS:
        extras = random.sample(pool, EXTRA_FROM_OTHER_SHEETS)
    else:
        # if no extra sheets or insufficient rows, sample from entire pool of sheets
        global_pool = []
        for sname in all_sheets:
            global_pool.extend(sheet_rows[sname])
        extras = random.sample(global_pool, EXTRA_FROM_OTHER_SHEETS)
    for e in extras:
        selected.append({'question': e['question'], 'A': e['A'], 'B': e['B'], 'C': e['C'], 'D': e['D'], 'correct': e['correct'], 'correct_raw': e['correct_raw']})

    if len(selected) != QUESTIONS_PER_VARIANT:
        print(f"Warning: variant {variant_num} has {len(selected)} questions (expected {QUESTIONS_PER_VARIANT})")

    # Shuffle overall question order
    random.shuffle(selected)

    # For each question shuffle options and determine new correct label
    final_questions = []
    answer_key = []
    for q in selected:
        shuffled_row, new_correct = shuffle_question(q)
        # ensure we have 4 option keys
        for lab in ['A','B','C','D']:
            if lab not in shuffled_row:
                shuffled_row[lab] = ''
        final_questions.append({'question': shuffled_row['question'], 'A': shuffled_row['A'], 'B': shuffled_row['B'], 'C': shuffled_row['C'], 'D': shuffled_row['D']})
        answer_key.append(new_correct if new_correct else '')

    # Generate random booklet id
    booklet_id = ''.join(random.choices(string.digits, k=BOOKLET_ID_LENGTH))

    # File names
    variant_fname = os.path.join(VARIANTS_FOLDER, f"Variant_{variant_num:03d}.docx")
    key_fname = os.path.join(KEYS_FOLDER, f"Variant_{variant_num:03d}_Answers.docx")

    # Write DOCX files
    write_variant_docx(variant_fname, booklet_id, "Matematika va informatika (kunduzgi)", "", final_questions)
    write_key_docx(key_fname, variant_num, answer_key)

    if variant_num % 10 == 0:
        print(f"Generated {variant_num}/{VARIANT_COUNT} variants")

print("All variants generated.")
print("Variants folder:", VARIANTS_FOLDER)
print("Keys folder:", KEYS_FOLDER)
